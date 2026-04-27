"""
Erstellt die DEMRE-Bedienungsanleitung als Word-Dokument.
Ausführen: python create_manual.py
"""
from docx import Document
from docx.shared import Pt, RGBColor, Cm, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import datetime

doc = Document()

# ── Seitenränder ──────────────────────────────────────────────────────────────
for section in doc.sections:
    section.left_margin   = Cm(2.5)
    section.right_margin  = Cm(2.5)
    section.top_margin    = Cm(2.5)
    section.bottom_margin = Cm(2.0)

# ── Hilfsfunktionen ───────────────────────────────────────────────────────────

def set_heading_style(para, level, text, color=None):
    para.clear()
    run = para.add_run(text)
    if level == 1:
        run.font.size = Pt(20)
        run.font.bold = True
        run.font.color.rgb = RGBColor(0x1A, 0x56, 0x9E)
    elif level == 2:
        run.font.size = Pt(14)
        run.font.bold = True
        run.font.color.rgb = RGBColor(0x1A, 0x56, 0x9E)
    elif level == 3:
        run.font.size = Pt(11)
        run.font.bold = True
        run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
    para.paragraph_format.space_before = Pt(12)
    para.paragraph_format.space_after  = Pt(4)


def add_h1(doc, text):
    p = doc.add_paragraph()
    set_heading_style(p, 1, text)
    return p


def add_h2(doc, text):
    p = doc.add_paragraph()
    set_heading_style(p, 2, text)
    return p


def add_h3(doc, text):
    p = doc.add_paragraph()
    set_heading_style(p, 3, text)
    return p


def add_body(doc, text):
    p = doc.add_paragraph(text)
    p.paragraph_format.space_after = Pt(6)
    p.runs[0].font.size = Pt(10) if p.runs else None
    return p


def add_bullet(doc, text, level=0):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.left_indent = Cm(0.5 + level * 0.5)
    p.paragraph_format.space_after = Pt(3)
    run = p.add_run(text)
    run.font.size = Pt(10)
    return p


def add_note(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent  = Cm(0.5)
    p.paragraph_format.right_indent = Cm(0.5)
    p.paragraph_format.space_after  = Pt(6)
    run = p.add_run("Hinweis:  ")
    run.font.bold = True
    run.font.color.rgb = RGBColor(0xC0, 0x50, 0x00)
    run.font.size = Pt(10)
    run2 = p.add_run(text)
    run2.font.size = Pt(10)
    run2.font.color.rgb = RGBColor(0x44, 0x44, 0x44)
    # leichter grauer Hintergrund via Shading ist per oxml möglich, aber optional
    return p


def add_step(doc, number, text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.5)
    p.paragraph_format.space_after = Pt(4)
    run1 = p.add_run(f"{number}.  ")
    run1.font.bold = True
    run1.font.size = Pt(10)
    run2 = p.add_run(text)
    run2.font.size = Pt(10)
    return p


def add_separator(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "4")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "AAAAAA")
    pBdr.append(bottom)
    pPr.append(pBdr)
    return p


# ── DECKBLATT ─────────────────────────────────────────────────────────────────

title_para = doc.add_paragraph()
title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
title_para.paragraph_format.space_before = Pt(60)
r = title_para.add_run("DEMRE")
r.font.size  = Pt(36)
r.font.bold  = True
r.font.color.rgb = RGBColor(0x1A, 0x56, 0x9E)

sub_para = doc.add_paragraph()
sub_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
r2 = sub_para.add_run("Rechnungs- und Vertragsverwaltung")
r2.font.size = Pt(18)
r2.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

doc.add_paragraph()
company_para = doc.add_paragraph()
company_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
r3 = company_para.add_run("Demme Immobilien Verwaltung GmbH")
r3.font.size = Pt(13)
r3.font.bold = True

version_para = doc.add_paragraph()
version_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
r4 = version_para.add_run(f"Bedienungsanleitung  •  Stand: {datetime.date.today().strftime('%d.%m.%Y')}")
r4.font.size  = Pt(11)
r4.font.color.rgb = RGBColor(0x77, 0x77, 0x77)

doc.add_page_break()

# ── INHALTSVERZEICHNIS (manuell) ──────────────────────────────────────────────
add_h1(doc, "Inhaltsverzeichnis")
toc_entries = [
    ("1", "Systemübersicht"),
    ("2", "Anmeldung"),
    ("3", "Dashboard"),
    ("4", "Kunden"),
    ("5", "Artikel (Leistungskatalog)"),
    ("6", "Abo-Rechnungen (Verträge)"),
    ("7", "Ausgangsrechnungen"),
    ("8", "Kreditoren"),
    ("9", "Eingangsrechnungen"),
    ("10", "Belege"),
    ("11", "Häufige Fragen & Lösungen"),
]
for num, title in toc_entries:
    p = doc.add_paragraph()
    r = p.add_run(f"{num}.  {title}")
    r.font.size = Pt(11)
    p.paragraph_format.space_after = Pt(4)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# 1  SYSTEMÜBERSICHT
# ══════════════════════════════════════════════════════════════════════════════
add_h1(doc, "1.  Systemübersicht")
add_body(doc,
    "DEMRE ist das interne Rechnungs- und Vertragsverwaltungssystem der "
    "Demme Immobilien Verwaltung GmbH. Es bildet den gesamten Abrechnungsprozess ab – "
    "von der Stammdatenpflege über die automatische Rechnungserstellung bis hin zur "
    "Verwaltung von Eingangsrechnungen und Belegen."
)

add_h2(doc, "Hauptfunktionen im Überblick")
features = [
    "Kundenverwaltung (Mieter, Eigentümer, Dienstleister)",
    "Artikelkatalog (Leistungen mit Preisen und Mehrwertsteuersätzen)",
    "Abo-Rechnungen / Verträge mit automatischem Rechnungslauf",
    "Ausgangsrechnungen im ZUGFeRD/Factur-X-Format (PDF + XML)",
    "Kreditorenverwaltung",
    "Eingangsrechnungen mit Datei-Upload",
    "Belegverwaltung",
]
for f in features:
    add_bullet(doc, f)

add_h2(doc, "Systemzugang")
add_bullet(doc, "Web-Browser (Chrome, Firefox, Edge empfohlen)")
add_bullet(doc, "Keine Installation erforderlich")
add_bullet(doc, "Zugang über das lokale Netzwerk oder den Produktionsserver")

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# 2  ANMELDUNG
# ══════════════════════════════════════════════════════════════════════════════
add_h1(doc, "2.  Anmeldung")
add_body(doc, "Beim Aufruf der Anwendung erscheint die Anmeldemaske.")

add_h2(doc, "Anmeldung durchführen")
add_step(doc, 1, "Browser öffnen und die Adresse der Anwendung aufrufen.")
add_step(doc, 2, "E-Mail-Adresse eingeben.")
add_step(doc, 3, "Passwort eingeben.")
add_step(doc, 4, 'Schaltfläche "Anmelden" klicken oder Enter drücken.')
add_step(doc, 5, "Bei erfolgreicher Anmeldung erscheint das Dashboard.")

add_note(doc,
    "Bei falschen Anmeldedaten erscheint eine Fehlermeldung. "
    "Wenden Sie sich bei Problemen an den Administrator."
)

add_h2(doc, "Abmelden")
add_body(doc,
    'Klicken Sie oben rechts auf das Benutzer-Symbol und wählen Sie "Abmelden". '
    "Die Sitzung wird beendet und Sie werden zur Anmeldeseite weitergeleitet."
)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# 3  DASHBOARD
# ══════════════════════════════════════════════════════════════════════════════
add_h1(doc, "3.  Dashboard")
add_body(doc,
    "Das Dashboard ist die Startseite nach der Anmeldung. "
    "Es gibt einen schnellen Überblick über den aktuellen Stand der Buchhaltung."
)

add_h2(doc, "Übersichtskacheln")
add_bullet(doc, "Offene Rechnungen – Anzahl und Gesamtbetrag unbezahlter Ausgangsrechnungen")
add_bullet(doc, "Aktive Verträge – Anzahl laufender Abo-Rechnungen")
add_bullet(doc, "Kunden – Gesamtanzahl angelegter Kunden")
add_bullet(doc, "Artikel – Anzahl aktiver Artikel im Leistungskatalog")

add_h2(doc, "Navigation")
add_body(doc,
    "Die Hauptnavigation befindet sich in der linken Seitenleiste. "
    "Durch Klick auf einen Menüpunkt gelangen Sie zum jeweiligen Bereich."
)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# 4  KUNDEN
# ══════════════════════════════════════════════════════════════════════════════
add_h1(doc, "4.  Kunden")
add_body(doc,
    "Im Bereich Kunden werden alle Empfänger von Ausgangsrechnungen verwaltet. "
    "Das können Mieter, Eigentümer, Gewerbetreibende oder andere Auftraggeber sein."
)

add_h2(doc, "Kundenliste")
add_bullet(doc, "Zeigt alle angelegten Kunden in einer Tabelle.")
add_bullet(doc, "Über das Suchfeld können Sie nach Name, E-Mail oder Kundennummer filtern.")
add_bullet(doc, "Ein Klick auf eine Zeile öffnet die Detailansicht.")

add_h2(doc, "Neuen Kunden anlegen")
add_step(doc, 1, 'Schaltfläche "Neuer Kunde" (oben rechts) klicken.')
add_step(doc, 2, "Pflichtfelder ausfüllen: Kundenname, Straße, PLZ, Ort, Land.")
add_step(doc, 3, "Optionale Felder: E-Mail, Telefon, USt-IdNr., Steuernummer.")
add_step(doc, 4, "Kundentyp wählen: Privatkunde oder Geschäftskunde.")
add_step(doc, 5, 'Schaltfläche "Speichern" klicken.')

add_note(doc,
    "Mit der TAB-Taste springen Sie bequem von Feld zu Feld. "
    "Pflichtfelder sind mit einem Sternchen (*) gekennzeichnet."
)

add_h2(doc, "Kunden bearbeiten")
add_step(doc, 1, "Kunden in der Liste auswählen (Klick auf die Zeile).")
add_step(doc, 2, 'Schaltfläche "Bearbeiten" (Stift-Symbol) klicken.')
add_step(doc, 3, "Felder anpassen.")
add_step(doc, 4, 'Schaltfläche "Speichern" klicken.')

add_h2(doc, "Kunden löschen")
add_body(doc,
    "Kunden können nur gelöscht werden, wenn ihnen keine Rechnungen oder Verträge zugeordnet sind. "
    "Andernfalls erscheint eine Fehlermeldung."
)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# 5  ARTIKEL
# ══════════════════════════════════════════════════════════════════════════════
add_h1(doc, "5.  Artikel (Leistungskatalog)")
add_body(doc,
    "Artikel sind die Grundbausteine aller Rechnungen. "
    "Jeder Artikel beschreibt eine abrechenbare Leistung mit Preis und Mehrwertsteuersatz."
)

add_h2(doc, "Artikelliste")
add_bullet(doc, "Zeigt alle aktiven und inaktiven Artikel.")
add_bullet(doc, "Spalten: Artikelname, Beschreibung, Preis (netto), MwSt.-Satz, Einheit, Status.")

add_h2(doc, "Neuen Artikel anlegen")
add_step(doc, 1, 'Schaltfläche "Neuer Artikel" klicken.')
add_step(doc, 2, "Artikelname und Beschreibung eingeben.")
add_step(doc, 3, "Nettopreis eingeben (Dezimalkomma oder -punkt möglich, z. B. 49,90 oder 49.90).")
add_step(doc, 4, "Mehrwertsteuersatz auswählen (Standard: 19 %).")
add_step(doc, 5, "Einheit wählen (z. B. Monat, Stück, Stunde).")
add_step(doc, 6, "Abrechnungszyklus wählen: monatlich oder jährlich.")
add_step(doc, 7, 'Schaltfläche "Speichern" klicken.')

add_h2(doc, "Artikel bearbeiten oder deaktivieren")
add_body(doc,
    "Klicken Sie in der Artikelliste auf das Stift-Symbol einer Zeile. "
    "Im Dialog können Sie alle Felder ändern und den Artikel über den Status-Schalter "
    "aktivieren oder deaktivieren. Inaktive Artikel werden im Rechnungslauf nicht berücksichtigt."
)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# 6  ABO-RECHNUNGEN (VERTRÄGE)
# ══════════════════════════════════════════════════════════════════════════════
add_h1(doc, "6.  Abo-Rechnungen (Verträge)")
add_body(doc,
    "Verträge sind wiederkehrende Abrechnungsvereinbarungen mit einem Kunden. "
    "Jeder Vertrag enthält eine oder mehrere Positionen (Artikel), "
    "die beim Rechnungslauf automatisch als Ausgangsrechnungen erstellt werden."
)

add_h2(doc, "Vertragsliste")
add_bullet(doc, "Zeigt alle Verträge mit Kunde, Vertragsnummer, Zeitraum und Status.")
add_bullet(doc, "Klick auf eine Zeile öffnet die Vertragsdetailseite.")

add_h2(doc, "Neuen Vertrag anlegen")
add_step(doc, 1, 'Schaltfläche "Neuer Vertrag" klicken.')
add_step(doc, 2, "Kunden aus der Dropdown-Liste auswählen.")
add_step(doc, 3, "Vertragsnummer vergeben (oder automatisch generieren lassen).")
add_step(doc, 4, "Startdatum eingeben. Enddatum ist optional.")
add_step(doc, 5, 'Schaltfläche "Speichern" klicken.')
add_step(doc, 6, "Auf der Detailseite: Vertragsposition(en) über das Plus-Symbol hinzufügen.")

add_h2(doc, "Vertrag bearbeiten")
add_body(doc,
    'Auf der Vertragsdetailseite finden Sie oben rechts die Schaltfläche "Bearbeiten" (Stift-Symbol). '
    "Damit können Sie Vertragsnummer, Beschreibung, Start- und Enddatum sowie den Status ändern."
)

add_h2(doc, "Vertragsposition hinzufügen")
add_step(doc, 1, 'Schaltfläche "Position hinzufügen" (Plus-Symbol) klicken.')
add_step(doc, 2, "Artikel aus der Dropdown-Liste wählen.")
add_step(doc, 3, "Menge eingeben.")
add_step(doc, 4, "Optional: individuellen Einzelpreis überschreiben.")
add_step(doc, 5, "Status: Aktiv oder Inaktiv.")
add_step(doc, 6, 'Schaltfläche "Speichern" klicken.')

add_h2(doc, "Vertragsposition bearbeiten")
add_body(doc,
    "Klicken Sie in der Positionstabelle auf das Stift-Symbol der jeweiligen Zeile. "
    "Sie können Menge, Preis und Status anpassen."
)

add_h2(doc, "Rechnungslauf durchführen")
add_step(doc, 1, 'Im Bereich "Ausgangsrechnungen" auf "Rechnungslauf" klicken.')
add_step(doc, 2, "Abrechnungszeitraum eingeben (Von / Bis).")
add_step(doc, 3, 'Option "Sofort ausstellen" aktivieren, wenn die Rechnungen direkt den Status "Ausgestellt" erhalten sollen.')
add_step(doc, 4, 'Schaltfläche "Rechnungslauf starten" klicken.')
add_step(doc, 5, "Das System erstellt automatisch eine Rechnung je aktiver Vertragsposition.")

add_note(doc,
    "Der Doppellaufschutz verhindert, dass dieselbe Position für denselben Zeitraum "
    "zweimal abgerechnet wird. Stornierte Rechnungen werden dabei nicht berücksichtigt."
)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# 7  AUSGANGSRECHNUNGEN
# ══════════════════════════════════════════════════════════════════════════════
add_h1(doc, "7.  Ausgangsrechnungen")
add_body(doc,
    "Ausgangsrechnungen werden entweder automatisch per Rechnungslauf erstellt "
    "oder manuell angelegt. Jede Rechnung kann als PDF (ZUGFeRD/Factur-X) heruntergeladen werden."
)

add_h2(doc, "Rechnungsstatus")
statuses = [
    ("Entwurf", "Rechnung wurde erstellt, aber noch nicht ausgestellt. Felder und Positionen können noch bearbeitet werden."),
    ("Ausgestellt", "Rechnung wurde an den Kunden übermittelt. Keine Änderungen mehr möglich."),
    ("Bezahlt", "Zahlung wurde verbucht."),
    ("Storniert", "Rechnung wurde storniert und ist für alle weiteren Berechnungen gesperrt."),
]
for status, desc in statuses:
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.5)
    p.paragraph_format.space_after = Pt(4)
    r1 = p.add_run(f"{status}: ")
    r1.font.bold = True
    r1.font.size = Pt(10)
    r2 = p.add_run(desc)
    r2.font.size = Pt(10)

add_h2(doc, "Rechnung manuell anlegen")
add_step(doc, 1, 'Schaltfläche "Neue Rechnung" klicken.')
add_step(doc, 2, "Kunden auswählen.")
add_step(doc, 3, "Rechnungsdatum und Fälligkeitsdatum eingeben.")
add_step(doc, 4, "Leistungszeitraum (Von / Bis) eingeben.")
add_step(doc, 5, "Positionen hinzufügen (Artikel, Menge, Einzelpreis).")
add_step(doc, 6, 'Schaltfläche "Speichern" klicken.')

add_h2(doc, "PDF herunterladen")
add_body(doc,
    'In der Rechnungsdetailansicht klicken Sie auf "PDF herunterladen". '
    "Das Dokument enthält das sichtbare PDF sowie eingebettete ZUGFeRD/Factur-X-XML-Daten."
)

add_h2(doc, "Rechnung ausstellen")
add_body(doc,
    'Klicken Sie in der Detailansicht auf "Ausstellen". '
    "Damit wechselt der Status von Entwurf auf Ausgestellt. "
    "Dieser Vorgang kann nicht rückgängig gemacht werden (außer durch Stornierung)."
)

add_h2(doc, "Rechnung stornieren")
add_body(doc,
    'Klicken Sie auf "Stornieren". Eine stornierte Rechnung bleibt im System gespeichert, '
    "wird aber bei allen Berechnungen (z. B. offene Posten, Rechnungslauf-Duplikatprüfung) "
    "nicht berücksichtigt."
)

add_h2(doc, "Rechnungspositionen bearbeiten (nur Entwurf)")
add_body(doc,
    "Positionen können nur bei Rechnungen im Status Entwurf geändert werden. "
    "Öffnen Sie die Detailansicht und nutzen Sie die Plus- bzw. Stift-Symbole in der Positionstabelle."
)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# 8  KREDITOREN
# ══════════════════════════════════════════════════════════════════════════════
add_h1(doc, "8.  Kreditoren")
add_body(doc,
    "Kreditoren sind Lieferanten und Dienstleister, von denen das Unternehmen Rechnungen erhält. "
    "Die Kreditorenverwaltung dient als Stammdatenbasis für Eingangsrechnungen."
)

add_h2(doc, "Kreditor anlegen")
add_step(doc, 1, 'Schaltfläche "Neuer Kreditor" klicken.')
add_step(doc, 2, "Name, Adresse und optional Bankdaten eingeben.")
add_step(doc, 3, 'Schaltfläche "Speichern" klicken.')

add_h2(doc, "Kreditor bearbeiten")
add_body(doc,
    "Klicken Sie auf das Stift-Symbol in der Kreditorenliste, um die Stammdaten zu aktualisieren."
)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# 9  EINGANGSRECHNUNGEN
# ══════════════════════════════════════════════════════════════════════════════
add_h1(doc, "9.  Eingangsrechnungen")
add_body(doc,
    "Im Bereich Eingangsrechnungen werden alle erhaltenen Rechnungen von Kreditoren erfasst. "
    "Rechnungen können manuell eingegeben oder als Datei hochgeladen werden."
)

add_h2(doc, "Eingangsrechnung anlegen")
add_step(doc, 1, 'Schaltfläche "Neue Eingangsrechnung" klicken.')
add_step(doc, 2, "Kreditor auswählen oder neu anlegen.")
add_step(doc, 3, "Rechnungsnummer, Rechnungsdatum und Betrag eingeben.")
add_step(doc, 4, "Optional: PDF oder andere Datei hochladen.")
add_step(doc, 5, 'Schaltfläche "Speichern" klicken.')

add_h2(doc, "Zahlungsstatus")
add_bullet(doc, "Offen – Rechnung ist erfasst, aber noch nicht bezahlt.")
add_bullet(doc, "Bezahlt – Zahlung wurde vermerkt.")
add_bullet(doc, "Storniert – Rechnung wurde storniert.")

add_h2(doc, "Automatischer Import (Überwachungsordner)")
add_body(doc,
    'Das System überwacht den Ordner "Eingangsrechnungen" auf dem Server. '
    "Abgelegte PDF-Dateien werden automatisch erkannt und als neue Eingangsrechnungen importiert. "
    "Die Original-Dateien bleiben im Ordner erhalten."
)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# 10  BELEGE
# ══════════════════════════════════════════════════════════════════════════════
add_h1(doc, "10.  Belege")
add_body(doc,
    "Im Belegbereich können allgemeine Ausgabenbelege (z. B. Kassenquittungen, Reisekostenabrechnungen) "
    "erfasst und Dateien hochgeladen werden."
)

add_h2(doc, "Beleg anlegen")
add_step(doc, 1, 'Schaltfläche "Neuer Beleg" klicken.')
add_step(doc, 2, "Datum, Beschreibung und Betrag eingeben.")
add_step(doc, 3, "Kategorie auswählen.")
add_step(doc, 4, "Datei (PDF, JPG, PNG) hochladen.")
add_step(doc, 5, 'Schaltfläche "Speichern" klicken.')

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# 11  HÄUFIGE FRAGEN & LÖSUNGEN
# ══════════════════════════════════════════════════════════════════════════════
add_h1(doc, "11.  Häufige Fragen & Lösungen")

faqs = [
    (
        "Ich kann eine Rechnung nicht bearbeiten.",
        "Rechnungen können nur im Status Entwurf bearbeitet werden. "
        "Sobald eine Rechnung ausgestellt oder bezahlt ist, sind keine Änderungen mehr möglich. "
        "Bei einem Fehler stornieren Sie die Rechnung und legen eine neue an."
    ),
    (
        "Der Rechnungslauf erzeugt keine neuen Rechnungen.",
        "Prüfen Sie folgende Punkte: (1) Sind die betroffenen Vertragspositionens aktiv? "
        "(2) Wurde für denselben Zeitraum bereits eine Rechnung erstellt (Doppellaufschutz)? "
        "(3) Ist das Enddatum des Vertrags abgelaufen?"
    ),
    (
        "Die PDF-Rechnung lässt sich nicht öffnen.",
        "Stellen Sie sicher, dass ein aktueller PDF-Reader installiert ist. "
        "Die Datei enthält eingebettete ZUGFeRD-XML-Daten – das ist korrekt und kein Fehler."
    ),
    (
        "Ich habe das Passwort vergessen.",
        "Wenden Sie sich an den Administrator, der Ihr Passwort zurücksetzen kann."
    ),
    (
        "Ein Artikel taucht im Rechnungslauf nicht auf.",
        "Überprüfen Sie, ob der Artikel in der Vertragsposition aktiv ist und ob der Artikel selbst "
        "nicht deaktiviert wurde."
    ),
    (
        "Die Anwendung zeigt eine Fehlermeldung (Fehler 500).",
        "Bitte notieren Sie die genaue Fehlermeldung und wenden Sie sich an den Administrator. "
        "Es handelt sich um einen serverseitigen Fehler, der in den Logs nachvollzogen werden kann."
    ),
    (
        "Kann ich einen Kunden löschen, der noch Rechnungen hat?",
        "Nein. Kunden mit zugeordneten Rechnungen oder Verträgen können nicht gelöscht werden. "
        "Deaktivieren Sie den Kunden stattdessen, um ihn aus Listen auszublenden."
    ),
]

for question, answer in faqs:
    add_h3(doc, question)
    add_body(doc, answer)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# SCHLUSSSEITE
# ══════════════════════════════════════════════════════════════════════════════
doc.add_paragraph()
doc.add_paragraph()
footer_para = doc.add_paragraph()
footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = footer_para.add_run("Demme Immobilien Verwaltung GmbH  •  DEMRE Bedienungsanleitung")
r.font.size = Pt(9)
r.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

date_para = doc.add_paragraph()
date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
r2 = date_para.add_run(f"Stand: {datetime.date.today().strftime('%d.%m.%Y')}")
r2.font.size = Pt(9)
r2.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

# ── Speichern ─────────────────────────────────────────────────────────────────
output_path = "DEMRE_Bedienungsanleitung.docx"
doc.save(output_path)
print(f"Dokument gespeichert: {output_path}")
